import os
import csv
import json
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def create_default_template():
    """Generates default_template.docx with Jinja2 placeholders."""
    doc = docx.Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("VULNERABILITY ASSESSMENT REPORT")
    run_title.font.name = 'Inter'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(99, 102, 241) # Indigo #6366f1
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Confidential Security Assessment Details")
    run_sub.font.name = 'Inter'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(107, 114, 128) # Grey #6b7280

    doc.add_paragraph() # Spacer

    # 1. Metadata Block (Table)
    doc.add_heading("1. Assessment Metadata", level=1)
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = 'Light Shading Accent 1'
    
    meta_data = [
        ("Client Name:", "{{ client_name }}"),
        ("Assessment Type:", "{{ scan_name }}"),
        ("Scan Date:", "{{ scan_date }}"),
        ("Target Scope:", "{{ target_scope }}")
    ]
    
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(val)

    doc.add_paragraph() # Spacer

    # 2. Executive Summary Metrics
    doc.add_heading("2. Executive Summary & Metrics", level=1)
    doc.add_paragraph("The security scan discovered vulnerabilities classified by severity below:")

    metrics_table = doc.add_table(rows=2, cols=5)
    metrics_table.style = 'Table Grid'
    
    # Headers
    headers = ["Critical", "High", "Medium", "Low", "Total Findings"]
    colors = ["EF4444", "F97316", "EAB308", "3B82F6", "6366F1"]
    
    hdr_row = metrics_table.rows[0]
    for idx, (h, col) in enumerate(zip(headers, colors)):
        cell = hdr_row.cells[idx]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, col)
        
    val_row = metrics_table.rows[1]
    metrics_placeholders = [
        "{{ metrics.critical }}",
        "{{ metrics.high }}",
        "{{ metrics.medium }}",
        "{{ metrics.low }}",
        "{{ metrics.total }}"
    ]
    for idx, p in enumerate(metrics_placeholders):
        cell = val_row.cells[idx]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(p)
        run.bold = True

    doc.add_paragraph() # Spacer
    doc.add_paragraph("Report compiled by: {{ tester_name }}")
    doc.add_paragraph() # Spacer

    # 3. Findings Table
    doc.add_heading("3. Detailed Vulnerability Findings", level=1)
    doc.add_paragraph("The following table lists the discovered vulnerability findings, hosts, description details, and remediation steps.")

    findings_table = doc.add_table(rows=4, cols=4)
    findings_table.style = 'Table Grid'
    
    # Row 0: Set headers
    hdr_cells = findings_table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run("Severity").bold = True
    hdr_cells[1].paragraphs[0].add_run("Vulnerability Name").bold = True
    hdr_cells[2].paragraphs[0].add_run("Target").bold = True
    hdr_cells[3].paragraphs[0].add_run("Description & Solution").bold = True
    
    # Style header cells background color (light indigo/slate)
    for cell in hdr_cells:
        set_cell_background(cell, "ECECFD")

    # Row 1: Start Loop Row
    findings_table.rows[1].cells[0].paragraphs[0].text = "{%tr for vuln in vulnerabilities %}"

    # Row 2: Data Row
    data_cells = findings_table.rows[2].cells
    
    p0 = data_cells[0].paragraphs[0]
    p0.add_run("{{ vuln.severity }}")
    
    p1 = data_cells[1].paragraphs[0]
    p1.add_run("{{ vuln.title }}").bold = True
    
    p2 = data_cells[2].paragraphs[0]
    p2.add_run("{{ vuln.host }}{% if vuln.port and vuln.port != 'N/A' %}:{{ vuln.port }}{% endif %}")
    
    p3 = data_cells[3].paragraphs[0]
    r_desc = p3.add_run("Description:\n")
    r_desc.bold = True
    p3.add_run("{{ vuln.description }}\n\n")
    r_fix = p3.add_run("Remediation:\n")
    r_fix.bold = True
    p3.add_run("{{ vuln.remediation }}")

    # Row 3: End Loop Row
    findings_table.rows[3].cells[0].paragraphs[0].text = "{%tr endfor %}"

    # Set column widths
    findings_table.columns[0].width = Inches(1.1)
    findings_table.columns[1].width = Inches(1.8)
    findings_table.columns[2].width = Inches(1.2)
    findings_table.columns[3].width = Inches(2.4)

    # Save
    tmpl_path = os.path.join(os.path.dirname(__file__), 'default_template.docx')
    doc.save(tmpl_path)
    print(f"Created default template: {tmpl_path}")

def create_sample_csv():
    """Generates sample_nessus_scan.csv for testing."""
    csv_path = os.path.join(os.path.dirname(__file__), 'sample_nessus_scan.csv')
    
    # Columns matching standard Nessus CSV export
    headers = ["Plugin ID", "CVE", "CVSS v2.0 Base Score", "Risk", "Host", "Protocol", "Port", "Name", "Synopsis", "Description", "Solution"]
    
    rows = [
        ["104631", "CVE-2017-0143", "9.3", "Critical", "192.168.1.45", "tcp", "445", 
         "MS17-010: Security Update for Microsoft Windows SMB Server (4013389)", 
         "The remote Windows host is affected by multiple vulnerabilities in SMBv1.", 
         "An remote attacker could execute arbitrary code on the target host by sending specially crafted SMBv1 packets.", 
         "Microsoft has released security updates to resolve this vulnerability. Disable SMBv1 and apply patches."],
        ["34220", "CVE-2008-4250", "10.0", "High", "192.168.1.12", "tcp", "445", 
         "Microsoft Windows Server Service Remote Code Execution (MS08-067)", 
         "The remote Windows host is vulnerable to a remote code execution vulnerability in the Server service.", 
         "A remote attacker can execute arbitrary code with SYSTEM privileges on the affected host by sending a crafted RPC request.", 
         "Apply the MS08-067 patch immediately."],
        ["51111", "CVE-2010-2568", "7.5", "Medium", "192.168.1.12", "tcp", "139", 
         "LNK File Automatic Remote Code Execution Vulnerability (MS10-046)", 
         "The remote host contains a vulnerability that allows code execution via a shortcut (.lnk) file.", 
         "Windows Shell fails to properly parse shortcut files, allowing code execution when the icon is rendered.", 
         "Install Microsoft update KB2286198."],
        ["11213", "N/A", "5.0", "Low", "192.168.1.100", "tcp", "80", 
         "HTTP Server Header Disclosure", 
         "The remote web server discloses its software version in header fields.", 
         "The web server headers contain details about the underlying software and OS versions, assisting target profiling.", 
         "Configure the web server to suppress the Server header (e.g. ServerTokens ProductOnly)."]
    ]
    
    with open(csv_path, mode='w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        writer.writerows(rows)
    print(f"Created sample CSV: {csv_path}")

def create_sample_json():
    """Generates sample_scan.json for testing."""
    json_path = os.path.join(os.path.dirname(__file__), 'sample_scan.json')
    
    data = {
        "metadata": {
            "scan_name": "API Service Security Audit",
            "scan_date": "2026-06-05",
            "target_scope": "api.prod.acme.corp",
            "client_name": "Acme Services LLC",
            "tester_name": "External Pentest Team"
        },
        "vulnerabilities": [
            {
                "title": "SQL Injection in Search API",
                "severity": "CRITICAL",
                "host": "api.prod.acme.corp",
                "port": "443",
                "description": "The 'query' parameter of the /api/v1/search endpoint does not properly sanitize user inputs, allowing SQL queries to be run directly on the database.",
                "remediation": "Implement parameterized queries or use a secure ORM framework to sanitize user inputs."
            },
            {
                "title": "Broken Object Level Authorization (BOLA)",
                "severity": "HIGH",
                "host": "api.prod.acme.corp",
                "port": "443",
                "description": "Users can access metadata of other accounts by modifying the 'account_id' parameter in /api/v1/accounts/{id} without active verification checks.",
                "remediation": "Verify authorization tokens against the requested object ownership before serving data."
            },
            {
                "title": "Weak JWT Signature Verification",
                "severity": "MEDIUM",
                "host": "api.prod.acme.corp",
                "port": "443",
                "description": "The authentication server accepts JWT tokens with the signature algorithm set to 'none', bypassing authentication checks.",
                "remediation": "Enforce strict signature validation using RS256/HS256 and reject tokens containing algorithm parameter set to 'none'."
            }
        ]
    }
    
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4)
    print(f"Created sample JSON: {json_path}")

def create_sample_text():
    """Generates sample_raw_scan.txt for testing."""
    txt_path = os.path.join(os.path.dirname(__file__), 'sample_raw_scan.txt')
    
    content = """=========================================
RAW VULNERABILITY SCANNER EXPORT
=========================================
Client: Acme Corp
Target Scope: 192.168.20.1-192.168.20.50
Date: 2026-06-08

[CRITICAL] Broken Access Control on File Server
Host: 192.168.20.10
Port: 80
Description: The backup folder /backup/db/ is accessible without password authentication, exposing raw SQL backups.
Remediation: Restrict HTTP folder listings and enforce active session verification on the directory.

[HIGH] Outdated Apache Apache Tomcat Vulnerability
Host: 192.168.20.15
Port: 8080
Description: The server runs Tomcat version 8.5.3 which has vulnerabilities to remote code execution.
Remediation: Upgrade the Apache Tomcat installation to version 8.5.50 or higher.

[MEDIUM] SSH Weak Ciphers Enabled
Host: 192.168.20.20
Port: 22
Description: The SSH server is configured to support weak MAC algorithms and CBC ciphers.
Remediation: Disable CBC mode ciphers and weak HMAC algorithms in sshd_config.
"""
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"Created sample TXT: {txt_path}")

if __name__ == '__main__':
    create_default_template()
    create_sample_csv()
    create_sample_json()
    create_sample_text()
